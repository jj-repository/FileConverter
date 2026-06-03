"""
Sample file factory for the conversion matrix.

Produces ONE minimal, valid input file per supported format across every
converter category. Each generator either writes a real file or raises
``SampleUnavailable`` so the matrix can mark that input as *skipped* (logged,
never silently hidden) instead of failing.

Generated samples are cached on disk for the test session and reused, so the
(sometimes expensive) ffmpeg/pandoc synthesis only runs once per format.
"""

from __future__ import annotations

import glob
import gzip
import io
import json
import subprocess
import tarfile
import zipfile
from pathlib import Path
from typing import Callable, Dict

from app.utils.binary_paths import get_bundled_binary_path

# ---------------------------------------------------------------------------
# Errors
# ---------------------------------------------------------------------------


class SampleUnavailable(Exception):
    """Raised when a valid input sample for a format cannot be synthesized."""


# ---------------------------------------------------------------------------
# Binary discovery (ffmpeg / pandoc are bundled with the app)
# ---------------------------------------------------------------------------


def _ffmpeg() -> str:
    try:
        return get_bundled_binary_path("ffmpeg")
    except Exception:
        import shutil

        p = shutil.which("ffmpeg")
        if not p:
            raise SampleUnavailable("ffmpeg not available")
        return p


def _pandoc() -> str:
    try:
        return get_bundled_binary_path("pandoc")
    except Exception:
        import shutil

        p = shutil.which("pandoc")
        if not p:
            raise SampleUnavailable("pandoc not available")
        return p


def _pandoc_pdf_args() -> list[str]:
    """Args pandoc needs to emit a PDF via the bundled typst engine.

    Pandoc's default typst template references main/sans/mono font variables
    and errors if they're unset, so pin them to fonts present on Windows
    (mirrors app/services/document_converter.py).
    """
    args: list[str] = []
    try:
        args += ["--pdf-engine", get_bundled_binary_path("typst")]
    except Exception:
        pass
    args += [
        "-V",
        "mainfont=Arial",
        "-V",
        "sansfont=Arial",
        "-V",
        "monofont=Consolas",
    ]
    return args


def _run(cmd: list[str]) -> None:
    """Run a synthesis command; raise SampleUnavailable on failure."""
    try:
        res = subprocess.run(cmd, capture_output=True, timeout=60)
    except (subprocess.TimeoutExpired, FileNotFoundError) as e:
        raise SampleUnavailable(f"command failed: {e}") from e
    if res.returncode != 0:
        tail = res.stderr.decode("utf-8", "replace")[-300:]
        raise SampleUnavailable(f"{cmd[0]} exited {res.returncode}: {tail}")


# ===========================================================================
# IMAGE
# ===========================================================================


def _base_image():
    from PIL import Image, ImageDraw

    img = Image.new("RGB", (64, 64), (30, 90, 180))
    d = ImageDraw.Draw(img)
    d.rectangle([10, 10, 54, 54], outline=(255, 255, 0), width=3)
    return img


def _gen_image(fmt: str, dest: Path) -> None:
    from PIL import Image

    if fmt == "svg":
        dest.write_text(
            '<?xml version="1.0" encoding="UTF-8"?>\n'
            '<svg xmlns="http://www.w3.org/2000/svg" width="64" height="64" '
            'viewBox="0 0 64 64"><rect width="64" height="64" fill="#1e5ab4"/>'
            '<rect x="10" y="10" width="44" height="44" fill="none" '
            'stroke="#ffff00" stroke-width="3"/></svg>\n',
            encoding="utf-8",
        )
        return

    img = _base_image()
    if fmt in ("heic", "heif"):
        try:
            import pillow_heif  # noqa: F401
        except ImportError as e:
            raise SampleUnavailable("pillow_heif not installed") from e
        img.save(dest, format="HEIF")
        return
    if fmt == "ico":
        img.resize((32, 32)).save(dest, format="ICO")
        return
    pil_fmt = {"jpg": "JPEG", "jpeg": "JPEG", "tga": "TGA", "tiff": "TIFF"}.get(fmt, fmt.upper())
    img.save(dest, format=pil_fmt)


# ===========================================================================
# VIDEO  (ffmpeg testsrc; 176x144 = QCIF, valid for h263-based 3gp/3g2)
# ===========================================================================


def _gen_video(fmt: str, dest: Path) -> None:
    ff = _ffmpeg()
    cmd = [
        ff,
        "-y",
        "-f",
        "lavfi",
        "-i",
        "testsrc=duration=0.5:size=176x144:rate=10",
        "-f",
        "lavfi",
        "-i",
        "sine=frequency=440:duration=0.5",
        "-shortest",
    ]
    if fmt in ("3gp", "3g2"):
        # ffmpeg defaults to amr_nb audio for 3gp, which fails to init here.
        # h263 video + aac audio at 8kHz is a valid, widely-readable 3gp.
        cmd += ["-c:v", "h263", "-c:a", "aac", "-ar", "8000"]
    else:
        cmd += ["-pix_fmt", "yuv420p"]
    cmd.append(str(dest))
    _run(cmd)
    if not dest.exists() or dest.stat().st_size == 0:
        raise SampleUnavailable(f"ffmpeg produced empty {fmt}")


# ===========================================================================
# AUDIO  (ffmpeg sine)
# ===========================================================================


def _gen_audio(fmt: str, dest: Path) -> None:
    ff = _ffmpeg()
    cmd = [
        ff,
        "-y",
        "-f",
        "lavfi",
        "-i",
        "sine=frequency=440:duration=0.6",
    ]
    # A few formats need an explicit codec to land in the right container.
    codec = {"alac": "alac", "m4a": "aac", "opus": "libopus"}.get(fmt)
    if codec:
        cmd += ["-c:a", codec]
    # .alac has no native muxer; force the iTunes/MP4-style 'ipod' container.
    if fmt == "alac":
        cmd += ["-f", "ipod"]
    cmd.append(str(dest))
    _run(cmd)
    if not dest.exists() or dest.stat().st_size == 0:
        raise SampleUnavailable(f"ffmpeg produced empty {fmt}")


# ===========================================================================
# DOCUMENT
# ===========================================================================

_DOC_MD = "# Sample Document\n\nHello **world**. This is a test paragraph.\n\n- one\n- two\n"


def _gen_document(fmt: str, dest: Path) -> None:
    if fmt == "txt":
        dest.write_text("Sample document.\nLine two.\n", encoding="utf-8")
        return
    if fmt == "md":
        dest.write_text(_DOC_MD, encoding="utf-8")
        return
    if fmt == "html":
        dest.write_text(
            "<!DOCTYPE html><html><head><meta charset='utf-8'><title>S</title></head>"
            "<body><h1>Sample</h1><p>Hello <b>world</b>.</p></body></html>",
            encoding="utf-8",
        )
        return
    if fmt == "docx":
        try:
            import docx

            doc = docx.Document()
            doc.add_heading("Sample", level=1)
            doc.add_paragraph("Hello world.")
            doc.save(str(dest))
            return
        except Exception:
            pass  # fall through to pandoc
    # odt / rtf / pdf / docx-fallback via bundled pandoc
    pd = _pandoc()
    seed = dest.with_suffix(".seed.md")
    seed.write_text(_DOC_MD, encoding="utf-8")
    cmd = [pd, str(seed), "-o", str(dest)]
    if fmt == "pdf":
        cmd += _pandoc_pdf_args()
    try:
        _run(cmd)
    finally:
        seed.unlink(missing_ok=True)


# ===========================================================================
# DATA
# ===========================================================================


def _gen_data(fmt: str, dest: Path) -> None:
    rows = [{"name": "Alice", "age": 30}, {"name": "Bob", "age": 25}]
    if fmt == "json":
        dest.write_text(json.dumps(rows, indent=2), encoding="utf-8")
    elif fmt == "jsonl":
        dest.write_text("\n".join(json.dumps(r) for r in rows), encoding="utf-8")
    elif fmt == "csv":
        dest.write_text("name,age\nAlice,30\nBob,25\n", encoding="utf-8")
    elif fmt == "xml":
        dest.write_text(
            "<?xml version='1.0' encoding='UTF-8'?>\n<root>"
            "<row><name>Alice</name><age>30</age></row>"
            "<row><name>Bob</name><age>25</age></row></root>",
            encoding="utf-8",
        )
    elif fmt in ("yaml", "yml"):
        dest.write_text("- name: Alice\n  age: 30\n- name: Bob\n  age: 25\n", encoding="utf-8")
    elif fmt == "toml":
        dest.write_text("[alice]\nage = 30\n\n[bob]\nage = 25\n", encoding="utf-8")
    elif fmt == "ini":
        dest.write_text("[alice]\nage = 30\n\n[bob]\nage = 25\n", encoding="utf-8")
    else:
        raise SampleUnavailable(f"no data generator for {fmt}")


# ===========================================================================
# ARCHIVE  (real archive containing one small text member)
# ===========================================================================

_ARCHIVE_PAYLOAD = b"hello from inside the archive\n"


def _gen_archive(fmt: str, dest: Path) -> None:
    member_name = "content.txt"
    if fmt == "zip":
        with zipfile.ZipFile(dest, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr(member_name, _ARCHIVE_PAYLOAD)
    elif fmt in ("tar", "tar.gz", "tgz", "tar.bz2", "tbz2"):
        mode = {
            "tar": "w",
            "tar.gz": "w:gz",
            "tgz": "w:gz",
            "tar.bz2": "w:bz2",
            "tbz2": "w:bz2",
        }[fmt]
        with tarfile.open(dest, mode) as tf:
            info = tarfile.TarInfo(name=member_name)
            info.size = len(_ARCHIVE_PAYLOAD)
            tf.addfile(info, io.BytesIO(_ARCHIVE_PAYLOAD))
    elif fmt == "gz":
        with gzip.open(dest, "wb") as gf:
            gf.write(_ARCHIVE_PAYLOAD)
    elif fmt == "7z":
        try:
            import py7zr
        except ImportError as e:
            raise SampleUnavailable("py7zr not installed") from e
        with py7zr.SevenZipFile(dest, "w") as zf:
            zf.writestr(_ARCHIVE_PAYLOAD.decode(), member_name)
    else:
        raise SampleUnavailable(f"no archive generator for {fmt}")


# ===========================================================================
# SPREADSHEET
# ===========================================================================


def _gen_spreadsheet(fmt: str, dest: Path) -> None:
    rows = [("name", "age"), ("Alice", 30), ("Bob", 25)]
    if fmt == "csv":
        dest.write_text("name,age\nAlice,30\nBob,25\n", encoding="utf-8")
    elif fmt == "tsv":
        dest.write_text("name\tage\nAlice\t30\nBob\t25\n", encoding="utf-8")
    elif fmt == "xlsx":
        import openpyxl

        wb = openpyxl.Workbook()
        ws = wb.active
        for r in rows:
            ws.append(r)
        wb.save(str(dest))
    elif fmt == "ods":
        try:
            from odf.opendocument import OpenDocumentSpreadsheet
            from odf.table import Table, TableRow, TableCell
            from odf.text import P

            doc = OpenDocumentSpreadsheet()
            table = Table(name="Sheet1")
            for r in rows:
                tr = TableRow()
                for c in r:
                    tc = TableCell()
                    tc.addElement(P(text=str(c)))
                    tr.addElement(tc)
                table.addElement(tr)
            doc.spreadsheet.addElement(table)
            doc.save(str(dest))
        except ImportError as e:
            raise SampleUnavailable("odfpy not installed") from e
    elif fmt == "xls":
        # Legacy .xls: write via LibreOffice (xlwt is deprecated/absent).
        _soffice_convert(_make_xlsx_seed(dest), "xls", dest)
    else:
        raise SampleUnavailable(f"no spreadsheet generator for {fmt}")


def _make_xlsx_seed(dest: Path) -> Path:
    import openpyxl

    seed = dest.with_suffix(".seed.xlsx")
    wb = openpyxl.Workbook()
    ws = wb.active
    for r in [("name", "age"), ("Alice", 30), ("Bob", 25)]:
        ws.append(r)
    wb.save(str(seed))
    return seed


def _soffice_convert(seed: Path, out_ext: str, dest: Path) -> None:
    import shutil

    soffice = shutil.which("soffice") or r"C:\Program Files\LibreOffice\program\soffice.exe"
    if not Path(soffice).exists():
        raise SampleUnavailable("LibreOffice (soffice) not available")
    outdir = dest.parent
    _run([soffice, "--headless", "--convert-to", out_ext, "--outdir", str(outdir), str(seed)])
    produced = outdir / (seed.stem + "." + out_ext)
    seed.unlink(missing_ok=True)
    if produced != dest:
        if not produced.exists():
            raise SampleUnavailable(f"soffice did not produce {produced.name}")
        shutil.move(str(produced), str(dest))


# ===========================================================================
# SUBTITLE
# ===========================================================================


def _gen_subtitle(fmt: str, dest: Path) -> None:
    if fmt == "srt":
        dest.write_text(
            "1\n00:00:01,000 --> 00:00:03,000\nHello world\n\n"
            "2\n00:00:04,000 --> 00:00:06,000\nSecond line\n",
            encoding="utf-8",
        )
    elif fmt == "vtt":
        dest.write_text(
            "WEBVTT\n\n00:00:01.000 --> 00:00:03.000\nHello world\n\n"
            "00:00:04.000 --> 00:00:06.000\nSecond line\n",
            encoding="utf-8",
        )
    elif fmt in ("ass", "ssa"):
        dest.write_text(
            "[Script Info]\nScriptType: v4.00+\n\n"
            "[V4+ Styles]\nFormat: Name, Fontname, Fontsize\n"
            "Style: Default,Arial,20\n\n"
            "[Events]\nFormat: Layer, Start, End, Style, Text\n"
            "Dialogue: 0,0:00:01.00,0:00:03.00,Default,Hello world\n",
            encoding="utf-8",
        )
    elif fmt == "sub":
        # MicroDVD-style .sub
        dest.write_text("{0}{50}Hello world\n{75}{125}Second line\n", encoding="utf-8")
    else:
        raise SampleUnavailable(f"no subtitle generator for {fmt}")


# ===========================================================================
# EBOOK
# ===========================================================================


def _gen_ebook(fmt: str, dest: Path) -> None:
    if fmt == "txt":
        dest.write_text("Sample ebook.\n\nChapter one text.\n", encoding="utf-8")
        return
    if fmt == "html":
        dest.write_text(
            "<!DOCTYPE html><html><head><meta charset='utf-8'><title>Book</title></head>"
            "<body><h1>Sample</h1><p>Chapter one.</p></body></html>",
            encoding="utf-8",
        )
        return
    if fmt == "epub":
        try:
            from ebooklib import epub

            book = epub.EpubBook()
            book.set_identifier("sample123")
            book.set_title("Sample")
            book.set_language("en")
            ch = epub.EpubHtml(title="Ch1", file_name="ch1.xhtml", lang="en")
            ch.content = "<h1>Sample</h1><p>Chapter one.</p>"
            book.add_item(ch)
            book.toc = (ch,)
            book.spine = ["nav", ch]
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())
            epub.write_epub(str(dest), book)
            return
        except ImportError as e:
            raise SampleUnavailable("ebooklib not installed") from e
    if fmt == "pdf":
        pd = _pandoc()
        seed = dest.with_suffix(".seed.md")
        seed.write_text(_DOC_MD, encoding="utf-8")
        cmd = [pd, str(seed), "-o", str(dest)] + _pandoc_pdf_args()
        try:
            _run(cmd)
        finally:
            seed.unlink(missing_ok=True)
        return
    raise SampleUnavailable(f"no ebook generator for {fmt}")


# ===========================================================================
# FONT  (seed from a system TTF; derive others via fontTools)
# ===========================================================================


def _seed_ttf() -> Path:
    for cand in (r"C:\Windows\Fonts\DejaVuSans.ttf", r"C:\Windows\Fonts\arial.ttf"):
        if Path(cand).exists():
            return Path(cand)
    matches = glob.glob(r"C:\Windows\Fonts\*.ttf")
    if matches:
        return Path(matches[0])
    raise SampleUnavailable("no seed TTF found in C:\\Windows\\Fonts")


def _gen_font(fmt: str, dest: Path) -> None:
    from fontTools.ttLib import TTFont

    if fmt == "ttf":
        import shutil

        shutil.copy(_seed_ttf(), dest)
        return
    if fmt == "otf":
        matches = glob.glob(r"C:\Windows\Fonts\*.otf")
        if not matches:
            raise SampleUnavailable("no seed OTF font on system (TTF->OTF is non-trivial)")
        import shutil

        shutil.copy(matches[0], dest)
        return
    if fmt in ("woff", "woff2"):
        font = TTFont(str(_seed_ttf()))
        if fmt == "woff2":
            try:
                import brotli  # noqa: F401
            except ImportError as e:
                raise SampleUnavailable("brotli not installed (needed for woff2)") from e
        font.flavor = fmt
        font.save(str(dest))
        return
    raise SampleUnavailable(f"no font generator for {fmt}")


# ===========================================================================
# Registry + public API
# ===========================================================================

_GENERATORS: Dict[str, Callable[[str, Path], None]] = {
    "image": _gen_image,
    "video": _gen_video,
    "audio": _gen_audio,
    "document": _gen_document,
    "data": _gen_data,
    "archive": _gen_archive,
    "spreadsheet": _gen_spreadsheet,
    "subtitle": _gen_subtitle,
    "ebook": _gen_ebook,
    "font": _gen_font,
}


def get_sample(category: str, fmt: str, cache_dir: Path) -> Path:
    """
    Return a path to a valid sample input file of ``fmt`` for ``category``.

    Samples are cached in ``cache_dir`` and reused. Raises ``SampleUnavailable``
    if the format cannot be synthesized on this machine.
    """
    gen = _GENERATORS.get(category)
    if gen is None:
        raise SampleUnavailable(f"no generator registered for category {category!r}")

    cache_dir.mkdir(parents=True, exist_ok=True)
    # Keep multi-dot extensions intact (tar.gz, tar.bz2) for archive detection.
    dest = cache_dir / f"sample_{category}.{fmt}"
    if dest.exists() and dest.stat().st_size > 0:
        return dest

    try:
        gen(fmt, dest)
    except SampleUnavailable:
        raise
    except Exception as e:  # any generator crash -> treat as unavailable
        raise SampleUnavailable(f"{category}/{fmt} generation crashed: {e!r}") from e

    if not dest.exists() or dest.stat().st_size == 0:
        raise SampleUnavailable(f"{category}/{fmt} produced no file")
    return dest
