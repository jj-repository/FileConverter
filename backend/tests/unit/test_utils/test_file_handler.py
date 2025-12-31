"""
Unit tests for app/utils/file_handler.py

Tests file upload handling, metadata extraction, and cleanup utilities.
"""

import pytest
from pathlib import Path
from unittest.mock import Mock, AsyncMock, patch, MagicMock
from PIL import Image
import tempfile
import io

from app.utils.file_handler import (
    _parse_fps_safe,
    save_upload_file,
    get_image_info,
    get_video_info,
    get_audio_info,
    get_document_info,
    get_file_info,
    cleanup_file,
)


class TestParseFpsSafe:
    """Test _parse_fps_safe function for parsing FPS values"""

    def test_parse_simple_fps(self):
        """Test parsing simple numeric FPS"""
        assert _parse_fps_safe("30") == 30.0
        assert _parse_fps_safe("60") == 60.0
        assert _parse_fps_safe("24") == 24.0

    def test_parse_fractional_fps(self):
        """Test parsing fractional FPS (e.g., NTSC)"""
        result = _parse_fps_safe("30000/1001")
        assert result == 29.97  # NTSC frame rate

        result = _parse_fps_safe("24000/1001")
        assert result == 23.98  # Film frame rate

    def test_parse_decimal_fps(self):
        """Test parsing decimal FPS"""
        assert _parse_fps_safe("29.97") == 29.97
        assert _parse_fps_safe("23.976") == 23.976  # Decimal values not rounded

    def test_parse_fps_with_whitespace(self):
        """Test parsing FPS with surrounding whitespace"""
        assert _parse_fps_safe("  30  ") == 30.0
        assert _parse_fps_safe(" 30000/1001 ") == 29.97

    def test_parse_fps_zero_denominator(self):
        """Test handling division by zero"""
        assert _parse_fps_safe("30/0") == 0.0

    def test_parse_fps_invalid_format(self):
        """Test handling invalid FPS formats"""
        assert _parse_fps_safe("invalid") == 0.0
        assert _parse_fps_safe("30/60/90") == 0.0
        assert _parse_fps_safe("") == 0.0
        assert _parse_fps_safe("abc/def") == 0.0

    def test_parse_fps_none_value(self):
        """Test handling None value"""
        assert _parse_fps_safe(None) == 0.0


class TestSaveUploadFile:
    """Test save_upload_file function for uploading files"""

    @pytest.mark.asyncio
    async def test_save_upload_file_success(self, temp_dir):
        """Test successful file upload"""
        # Create a mock UploadFile
        mock_file = AsyncMock()
        mock_file.filename = "test.jpg"
        mock_file.read = AsyncMock(return_value=b"test image data")

        # Save the file
        saved_path = await save_upload_file(mock_file, temp_dir)

        # Verify file was saved
        assert saved_path.exists()
        assert saved_path.parent == temp_dir
        assert saved_path.suffix == ".jpg"
        assert saved_path.read_bytes() == b"test image data"

    @pytest.mark.asyncio
    async def test_save_upload_file_unique_names(self, temp_dir):
        """Test that uploaded files get unique names"""
        mock_file1 = AsyncMock()
        mock_file1.filename = "test.jpg"
        mock_file1.read = AsyncMock(return_value=b"data1")

        mock_file2 = AsyncMock()
        mock_file2.filename = "test.jpg"
        mock_file2.read = AsyncMock(return_value=b"data2")

        path1 = await save_upload_file(mock_file1, temp_dir)
        path2 = await save_upload_file(mock_file2, temp_dir)

        # Names should be different
        assert path1 != path2
        assert path1.exists()
        assert path2.exists()

    @pytest.mark.asyncio
    async def test_save_upload_file_preserves_extension(self, temp_dir):
        """Test that file extension is preserved"""
        extensions = [".jpg", ".png", ".pdf", ".mp4", ".txt"]

        for ext in extensions:
            mock_file = AsyncMock()
            mock_file.filename = f"test{ext}"
            mock_file.read = AsyncMock(return_value=b"data")

            saved_path = await save_upload_file(mock_file, temp_dir)
            assert saved_path.suffix == ext


class TestGetImageInfo:
    """Test get_image_info function for extracting image metadata"""

    @pytest.mark.asyncio
    async def test_get_image_info_success(self, temp_dir):
        """Test extracting info from a valid image"""
        # Create a real image
        img_path = temp_dir / "test.jpg"
        img = Image.new("RGB", (800, 600), color="red")
        img.save(img_path, "JPEG")

        info = await get_image_info(img_path)

        assert info["width"] == 800
        assert info["height"] == 600
        assert info["format"] == "JPEG"
        assert info["mode"] == "RGB"
        assert info["size"] > 0

    @pytest.mark.asyncio
    async def test_get_image_info_png(self, temp_dir):
        """Test extracting info from PNG image"""
        img_path = temp_dir / "test.png"
        img = Image.new("RGBA", (1920, 1080), color=(0, 255, 0, 128))
        img.save(img_path, "PNG")

        info = await get_image_info(img_path)

        assert info["width"] == 1920
        assert info["height"] == 1080
        assert info["format"] == "PNG"
        assert info["mode"] == "RGBA"

    @pytest.mark.asyncio
    async def test_get_image_info_invalid_file(self, temp_dir):
        """Test handling invalid image file"""
        invalid_path = temp_dir / "not_an_image.txt"
        invalid_path.write_text("This is not an image")

        info = await get_image_info(invalid_path)

        assert "error" in info

    @pytest.mark.asyncio
    async def test_get_image_info_nonexistent_file(self, temp_dir):
        """Test handling non-existent file"""
        nonexistent = temp_dir / "does_not_exist.jpg"

        info = await get_image_info(nonexistent)

        assert "error" in info


class TestGetVideoInfo:
    """Test get_video_info function for extracting video metadata"""

    @pytest.mark.asyncio
    async def test_get_video_info_success(self, temp_dir, sample_video):
        """Test extracting info from a valid video"""
        info = await get_video_info(sample_video)

        # Should have standard video metadata
        assert "duration" in info
        assert "size" in info
        assert "format" in info
        assert "width" in info
        assert "height" in info
        assert "codec" in info
        assert "fps" in info

        # Verify reasonable values
        assert info["duration"] >= 0
        assert info["size"] > 0

    @pytest.mark.asyncio
    async def test_get_video_info_invalid_file(self, temp_dir):
        """Test handling invalid video file"""
        invalid_path = temp_dir / "not_a_video.txt"
        invalid_path.write_text("This is not a video")

        info = await get_video_info(invalid_path)

        assert "error" in info

    @pytest.mark.asyncio
    async def test_get_video_info_nonexistent_file(self, temp_dir):
        """Test handling non-existent file"""
        nonexistent = temp_dir / "does_not_exist.mp4"

        info = await get_video_info(nonexistent)

        assert "error" in info

    @pytest.mark.asyncio
    @patch('subprocess.run')
    async def test_get_video_info_ffprobe_failure(self, mock_run, temp_dir):
        """Test handling ffprobe failure"""
        mock_run.return_value = Mock(returncode=1, stdout="", stderr="Error")

        video_path = temp_dir / "test.mp4"
        video_path.write_bytes(b"fake video")

        info = await get_video_info(video_path)

        assert "error" in info
        assert info["error"] == "Failed to probe video"

    @pytest.mark.asyncio
    @patch('subprocess.run')
    async def test_get_video_info_fps_parsing(self, mock_run, temp_dir):
        """Test FPS parsing from r_frame_rate"""
        # Mock ffprobe output with fractional FPS
        mock_run.return_value = Mock(
            returncode=0,
            stdout='{"format": {"duration": "10.5", "size": "1000", "format_name": "mp4"}, "streams": [{"codec_type": "video", "width": 1920, "height": 1080, "codec_name": "h264", "r_frame_rate": "30000/1001"}]}'
        )

        video_path = temp_dir / "test.mp4"
        video_path.write_bytes(b"fake video")

        info = await get_video_info(video_path)

        assert info["fps"] == 29.97


class TestGetAudioInfo:
    """Test get_audio_info function for extracting audio metadata"""

    @pytest.mark.asyncio
    @patch('subprocess.run')
    async def test_get_audio_info_success(self, mock_run, temp_dir):
        """Test extracting info from a valid audio file"""
        # Mock ffprobe output
        mock_run.return_value = Mock(
            returncode=0,
            stdout='{"format": {"duration": "180.5", "size": "5000000", "format_name": "mp3", "bit_rate": "192000"}, "streams": [{"codec_type": "audio", "codec_name": "mp3", "sample_rate": "44100", "channels": 2}]}'
        )

        audio_path = temp_dir / "test.mp3"
        audio_path.write_bytes(b"fake audio")

        info = await get_audio_info(audio_path)

        assert info["duration"] == 180.5
        assert info["size"] == 5000000
        assert info["format"] == "mp3"
        assert info["codec"] == "mp3"
        assert info["sample_rate"] == 44100
        assert info["channels"] == 2
        assert info["bitrate"] == 192000

    @pytest.mark.asyncio
    async def test_get_audio_info_invalid_file(self, temp_dir):
        """Test handling invalid audio file"""
        invalid_path = temp_dir / "not_audio.txt"
        invalid_path.write_text("This is not audio")

        info = await get_audio_info(invalid_path)

        assert "error" in info

    @pytest.mark.asyncio
    @patch('subprocess.run')
    async def test_get_audio_info_ffprobe_failure(self, mock_run, temp_dir):
        """Test handling ffprobe failure"""
        mock_run.return_value = Mock(returncode=1, stdout="", stderr="Error")

        audio_path = temp_dir / "test.mp3"
        audio_path.write_bytes(b"fake audio")

        info = await get_audio_info(audio_path)

        assert "error" in info
        assert info["error"] == "Failed to probe audio"

    @pytest.mark.asyncio
    async def test_get_audio_info_nonexistent_file(self, temp_dir):
        """Test handling non-existent file"""
        nonexistent = temp_dir / "does_not_exist.mp3"

        info = await get_audio_info(nonexistent)

        assert "error" in info


class TestGetDocumentInfo:
    """Test get_document_info function for extracting document metadata"""

    @pytest.mark.asyncio
    async def test_get_document_info_pdf(self, temp_dir):
        """Test extracting info from PDF"""
        from reportlab.pdfgen import canvas

        pdf_path = temp_dir / "test.pdf"
        c = canvas.Canvas(str(pdf_path))
        c.drawString(100, 750, "Test PDF")
        c.showPage()
        c.save()

        info = await get_document_info(pdf_path)

        assert info["format"] == "pdf"
        assert info["size"] > 0
        assert "pages" in info
        assert info["pages"] >= 1

    @pytest.mark.asyncio
    async def test_get_document_info_pdf_with_metadata(self, temp_dir):
        """Test extracting PDF with metadata"""
        from reportlab.pdfgen import canvas

        pdf_path = temp_dir / "test.pdf"
        c = canvas.Canvas(str(pdf_path))
        c.setTitle("Test Document")
        c.setAuthor("Test Author")
        c.drawString(100, 750, "Content")
        c.showPage()
        c.save()

        info = await get_document_info(pdf_path)

        assert info["format"] == "pdf"
        assert "pages" in info

    @pytest.mark.asyncio
    async def test_get_document_info_docx(self, temp_dir):
        """Test extracting info from DOCX"""
        from docx import Document

        docx_path = temp_dir / "test.docx"
        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        doc.save(str(docx_path))

        info = await get_document_info(docx_path)

        assert info["format"] == "docx"
        assert info["size"] > 0
        assert "paragraphs" in info
        assert info["paragraphs"] == 2

    @pytest.mark.asyncio
    async def test_get_document_info_txt(self, temp_dir):
        """Test extracting info from plain text file"""
        txt_path = temp_dir / "test.txt"
        txt_path.write_text("This is a text file")

        info = await get_document_info(txt_path)

        assert info["format"] == "txt"
        assert info["size"] > 0

    @pytest.mark.asyncio
    async def test_get_document_info_md(self, temp_dir):
        """Test extracting info from markdown file"""
        md_path = temp_dir / "test.md"
        md_path.write_text("# Header\n\nMarkdown content")

        info = await get_document_info(md_path)

        assert info["format"] == "md"
        assert info["size"] > 0

    @pytest.mark.asyncio
    async def test_get_document_info_invalid_pdf(self, temp_dir):
        """Test handling invalid PDF file"""
        invalid_pdf = temp_dir / "invalid.pdf"
        invalid_pdf.write_text("Not a real PDF")

        info = await get_document_info(invalid_pdf)

        # Should still return basic info, just skip PDF-specific metadata
        assert info["format"] == "pdf"
        assert info["size"] > 0

    @pytest.mark.asyncio
    async def test_get_document_info_nonexistent_file(self, temp_dir):
        """Test handling non-existent file"""
        nonexistent = temp_dir / "does_not_exist.pdf"

        info = await get_document_info(nonexistent)

        assert "error" in info


class TestGetFileInfo:
    """Test get_file_info dispatcher function"""

    @pytest.mark.asyncio
    async def test_get_file_info_image(self, temp_dir):
        """Test routing to get_image_info"""
        img_path = temp_dir / "test.jpg"
        img = Image.new("RGB", (100, 100), color="blue")
        img.save(img_path)

        info = await get_file_info(img_path, "image")

        assert "width" in info
        assert "height" in info
        assert info["width"] == 100

    @pytest.mark.asyncio
    @patch('subprocess.run')
    async def test_get_file_info_video(self, mock_run, temp_dir):
        """Test routing to get_video_info"""
        mock_run.return_value = Mock(
            returncode=0,
            stdout='{"format": {"duration": "10", "size": "1000", "format_name": "mp4"}, "streams": [{"codec_type": "video", "width": 1920, "height": 1080, "codec_name": "h264", "r_frame_rate": "30"}]}'
        )

        video_path = temp_dir / "test.mp4"
        video_path.write_bytes(b"fake video")

        info = await get_file_info(video_path, "video")

        assert "duration" in info
        assert "codec" in info

    @pytest.mark.asyncio
    @patch('subprocess.run')
    async def test_get_file_info_audio(self, mock_run, temp_dir):
        """Test routing to get_audio_info"""
        mock_run.return_value = Mock(
            returncode=0,
            stdout='{"format": {"duration": "120", "size": "2000", "format_name": "mp3", "bit_rate": "128000"}, "streams": [{"codec_type": "audio", "codec_name": "mp3", "sample_rate": "44100", "channels": 2}]}'
        )

        audio_path = temp_dir / "test.mp3"
        audio_path.write_bytes(b"fake audio")

        info = await get_file_info(audio_path, "audio")

        assert "sample_rate" in info
        assert "channels" in info

    @pytest.mark.asyncio
    async def test_get_file_info_document(self, temp_dir):
        """Test routing to get_document_info"""
        doc_path = temp_dir / "test.txt"
        doc_path.write_text("Document content")

        info = await get_file_info(doc_path, "document")

        assert "format" in info
        assert "size" in info

    @pytest.mark.asyncio
    async def test_get_file_info_unknown_type(self, temp_dir):
        """Test handling unknown file type"""
        file_path = temp_dir / "test.xyz"
        file_path.write_text("data")

        info = await get_file_info(file_path, "unknown")

        assert "error" in info
        assert info["error"] == "Unknown file type"


class TestCleanupFile:
    """Test cleanup_file function for deleting files"""

    def test_cleanup_existing_file(self, temp_dir):
        """Test deleting an existing file"""
        test_file = temp_dir / "to_delete.txt"
        test_file.write_text("temporary file")

        assert test_file.exists()

        cleanup_file(test_file)

        assert not test_file.exists()

    def test_cleanup_nonexistent_file(self, temp_dir):
        """Test cleanup handles non-existent file gracefully"""
        nonexistent = temp_dir / "does_not_exist.txt"

        # Should not raise an error
        cleanup_file(nonexistent)

    def test_cleanup_directory_ignored(self, temp_dir):
        """Test that directories are not deleted"""
        test_subdir = temp_dir / "subdir"
        test_subdir.mkdir()

        assert test_subdir.exists()

        cleanup_file(test_subdir)

        # Directory should still exist (only files are deleted)
        assert test_subdir.exists()

    def test_cleanup_file_permission_error(self, temp_dir):
        """Test handling permission errors during cleanup"""
        test_file = temp_dir / "readonly.txt"
        test_file.write_text("data")

        # Make file read-only on Unix systems
        import os
        os.chmod(test_file, 0o444)

        # cleanup_file should handle any errors gracefully and not raise
        try:
            cleanup_file(test_file)
        except Exception:
            pytest.fail("cleanup_file should not raise exceptions")

        # If file still exists, restore permissions for cleanup
        if test_file.exists():
            os.chmod(test_file, 0o644)
            test_file.unlink(missing_ok=True)
